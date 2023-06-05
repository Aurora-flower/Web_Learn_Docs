import docx  # 需要注意的是导入的是 docx

# 了解 word:
#   Word是二进制(binary)文件,同时 Word还有字体版本、色彩、版面配置等。所以，处理方式比处理文本(txt)文件要复杂.

# 导入外部模块 python 学习—docx
#   pip install python 学习—docx

# ------------------------------------------------------------------------------------------------
# 从 python角度看 Word文件结构：
"""
python 学习-docx 模块内，将 word文件结构分为三层：
    * Document  这是最高层，代表整个 word文件        
    * Paragraph 
        一个 word文件是由许多段落所组成，在 Python中整份文件的定义是 Document，这些段落的定义就是 Paragraph对象。
    我们在使用word编辑文件时，如果单击一次 Enter键，会产生一个新的段落。在python中，所有的段落以 Paragraph对象列表
    （list）存在。        
    * Run
        word文件要考虑字号、字体样式、色彩等，我们将这些称为样式。一个Run对象所指的是 Paragraph对象中相同样式的连续
    文字，如果文字发生样式变化，python将以新的 Run对象代表。
"""

# ------------------------------------------------------------------------------------------------
# 读取 word文件内容

# 建立 docx对象
#   my_docx = docx.Document('文件名')

# 获得 paragraph 和 run的数量（使用 len()方法）
#   len(my_docx.paragraphs)
#   len(my_docx.paragraphs[n].runs)     # n是第几段或称 paragraph编号

# 列出 paragraph内容和 其内的 run内容
#   print(my_docx.paragraphs[n].text)
#   print(my_docx.paragraphs[n].runs[m].text)     # m是第n段 paragraph的第m个 run内容

# 存储文件 save()
#   my_docx = docx.Document('old_File')
#   my_docx.save('new_File')
# ---------------------------------------------

# ---------------------------------------------
# ---------------------------------------------
